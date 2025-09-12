"""
RAG (Retrieval-Augmented Generation) service
"""
import uuid
import asyncio
from typing import List, Dict, Any, Optional, AsyncGenerator
from llama_index.core import Document as LlamaDocument
from llama_index.core.node_parser import SentenceSplitter
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.llms.openai import OpenAI
from llama_index.vector_stores.qdrant import QdrantVectorStore
from llama_index.core import VectorStoreIndex, StorageContext
from llama_index.core.retrievers import VectorIndexRetriever
from llama_index.core.query_engine import RetrieverQueryEngine
# from llama_index.core.llms.types import ChatMessage
from qdrant_client.http import models
import openai
from config import settings
from models import QueryResponse, Chunk

class RAGService:
    def __init__(self, qdrant_client):
        self.qdrant = qdrant_client
        self.embedding_model = self._get_embedding_model()
        self.llm = self._get_llm()
        self.vector_store = self._get_vector_store()
        self.index = self._get_or_create_index()
    
    def _get_embedding_model(self):
        """Get embedding model"""
        if settings.USE_LOCAL_EMBEDDINGS:
            from sentence_transformers import SentenceTransformer
            return SentenceTransformer(settings.LOCAL_EMBEDDING_MODEL)
        else:
            return OpenAIEmbedding(
                model=settings.EMBEDDING_MODEL,
                api_key=settings.OPENAI_API_KEY
            )
    
    def _get_llm(self):
        """Get LLM"""
        if settings.USE_LOCAL_LLM:
            from llama_index.llms.ollama import Ollama
            return Ollama(
                model=settings.LOCAL_LLM_MODEL,
                base_url=settings.OLLAMA_BASE_URL
            )
        else:
            return OpenAI(
                model=settings.LLM_MODEL,
                api_key=settings.OPENAI_API_KEY
            )
    
    def _get_vector_store(self):
        """Get Qdrant vector store"""
        return QdrantVectorStore(
            client=self.qdrant,
            collection_name="chunks"
        )
    
    def _get_or_create_index(self):
        """Get or create vector index"""
        storage_context = StorageContext.from_defaults(vector_store=self.vector_store)
        return VectorStoreIndex.from_vector_store(
            vector_store=self.vector_store,
            storage_context=storage_context,
            embed_model=self.embedding_model
        )
    
    async def ingest_document(self, document_id: str, file_path: str, content_type: str) -> bool:
        """Ingest a document into the vector store"""
        try:
            # Parse document based on content type
            text_content = await self._parse_document(file_path, content_type)
            if not text_content:
                return False
            
            # Create LlamaIndex document
            llama_doc = LlamaDocument(
                text=text_content,
                metadata={
                    "document_id": document_id,
                    "filename": file_path.split("/")[-1],
                    "content_type": content_type
                }
            )
            
            # Chunk the document
            chunk_size = settings.CHUNK_SIZE
            chunk_overlap = settings.CHUNK_OVERLAP
            
            splitter = SentenceSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap
            )
            
            nodes = splitter.get_nodes_from_documents([llama_doc])
            
            # Add metadata to nodes
            for i, node in enumerate(nodes):
                node.metadata.update({
                    "chunk_id": str(uuid.uuid4()),
                    "chunk_index": i,
                    "document_id": document_id,
                    "embedding_version": settings.EMBEDDING_VERSION
                })
            
            # Add to vector store
            self.index.insert_nodes(nodes)
            
            return True
        except Exception as e:
            print(f"Error ingesting document: {e}")
            return False
    
    async def _parse_document(self, file_path: str, content_type: str) -> Optional[str]:
        """Parse document based on content type"""
        try:
            if content_type == "application/pdf":
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                return text
            
            elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                from docx import Document
                doc = Document(file_path)
                return "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            elif content_type == "text/plain":
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            
            elif content_type == "text/html":
                from bs4 import BeautifulSoup
                with open(file_path, "r", encoding="utf-8") as f:
                    soup = BeautifulSoup(f.read(), "html.parser")
                    return soup.get_text()
            
            elif content_type == "text/markdown":
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            
            else:
                # Try to read as text
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
        
        except Exception as e:
            print(f"Error parsing document {file_path}: {e}")
            return None
    
    async def query(self, query_text: str, filters: Dict[str, Any] = None) -> QueryResponse:
        """Query the RAG system"""
        try:
            # Create retriever
            retriever = VectorIndexRetriever(
                index=self.index,
                similarity_top_k=5
            )
            
            # Create query engine
            query_engine = RetrieverQueryEngine.from_args(
                retriever=retriever,
                llm=self.llm
            )
            
            # Execute query
            response = query_engine.query(query_text)
            
            # Extract sources
            sources = []
            for node in response.source_nodes:
                sources.append({
                    "text": node.text[:200] + "..." if len(node.text) > 200 else node.text,
                    "score": node.score,
                    "metadata": node.metadata
                })
            
            return QueryResponse(
                answer=str(response),
                sources=sources,
                metadata={
                    "query": query_text,
                    "num_sources": len(sources)
                }
            )
        
        except Exception as e:
            print(f"Error in RAG query: {e}")
            return QueryResponse(
                answer=f"Error processing query: {str(e)}",
                sources=[],
                metadata={"error": str(e)}
            )
    
    async def query_stream(self, query_text: str, filters: Dict[str, Any] = None) -> AsyncGenerator[Dict[str, Any], None]:
        """Stream query response"""
        try:
            # Create retriever
            retriever = VectorIndexRetriever(
                index=self.index,
                similarity_top_k=5
            )
            
            # Create query engine with streaming
            query_engine = RetrieverQueryEngine.from_args(
                retriever=retriever,
                llm=self.llm
            )
            
            # Execute query
            response = query_engine.query(query_text)
            
            # Return the response
            yield {"type": "answer", "content": str(response)}
            
            # Add sources
            sources = []
            for node in response.source_nodes:
                sources.append({
                    "text": node.text[:200] + "..." if len(node.text) > 200 else node.text,
                    "score": node.score,
                    "metadata": node.metadata
                })
            
            yield {"type": "sources", "content": sources}
        
        except Exception as e:
            yield {"type": "error", "content": str(e)}
    
    async def delete_document_chunks(self, document_id: str) -> bool:
        """Delete all chunks for a document"""
        try:
            # Delete from Qdrant
            self.qdrant.delete(
                collection_name="chunks",
                points_selector=models.FilterSelector(
                    filter=models.Filter(
                        must=[
                            models.FieldCondition(
                                key="document_id",
                                match=models.MatchValue(value=document_id)
                            )
                        ]
                    )
                )
            )
            return True
        except Exception as e:
            print(f"Error deleting document chunks: {e}")
            return False
